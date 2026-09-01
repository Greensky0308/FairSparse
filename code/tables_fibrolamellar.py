"""Fibrolamellar-spectrum cohort table (xlsx + docx three-line table), read from cohort.tsv."""
import os, json
import pandas as pd
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA = os.path.join(os.path.dirname(__file__), "..", "data", "fibrolamellar")
OUT = os.path.join(os.path.dirname(__file__), "..", "tables")

def set_border(cell, edge, sz=6):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    b = tcPr.find(qn("w:tcBorders"))
    if b is None:
        b = OxmlElement("w:tcBorders"); tcPr.append(b)
    el = OxmlElement(f"w:{edge}"); el.set(qn("w:val"),"single"); el.set(qn("w:sz"),str(sz)); el.set(qn("w:space"),"0")
    b.append(el)

def three_line(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Table Grid"; t.autofit = True
    for i,h in enumerate(header):
        c = t.rows[0].cells[i]; c.text = str(h)
        for p in c.paragraphs:
            p.alignment=1
            for r in p.runs: r.font.size=Pt(9); r.font.bold=True
        set_border(c,"top",12); set_border(c,"bottom",6)
    for row in rows:
        cells = t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text = str(v)
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.size=Pt(9)
    for c in t.rows[-1].cells: set_border(c,"bottom",12)
    tblPr = t._tbl.tblPr; borders = OxmlElement("w:tblBorders")
    for edge in ("left","right","insideH","insideV"):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:val"),"none"); borders.append(e)
    tblPr.append(borders)

def write_both(name, header, rows):
    wb = Workbook(); ws = wb.active; ws.title="Sheet1"; ws.append(header)
    for r in rows: ws.append(list(r))
    wb.save(os.path.join(OUT, name+".xlsx"))
    doc = Document(); doc.add_heading(name.replace("_"," "), level=1); three_line(doc, header, rows)
    doc.save(os.path.join(OUT, name+".docx")); print("saved", name)

def table1():
    df = pd.read_csv(os.path.join(DATA, "cohort.tsv"), sep="\t")
    df["AGE"] = pd.to_numeric(df["AGE"])
    pos = df[df.fusion_status=="positive"]; neg = df[df.fusion_status=="negative"]
    header = ["Characteristic", "Value"]
    rows = [
        ["Patients, n", len(df)],
        ["Molecular subtype", ""],
        ["  Fusion-positive (FLC), n (%)", f"{len(pos)} ({len(pos)/len(df)*100:.0f}%)"],
        ["  Fusion-negative (BAP1), n (%)", f"{len(neg)} ({len(neg)/len(df)*100:.0f}%)"],
        ["Age (years), median (IQR)", ""],
        ["  Fusion-positive", f"{pos.AGE.median():.0f} ({pos.AGE.quantile(0.25):.0f}-{pos.AGE.quantile(0.75):.0f})"],
        ["  Fusion-negative", f"{neg.AGE.median():.0f} ({neg.AGE.quantile(0.25):.0f}-{neg.AGE.quantile(0.75):.0f})"],
        ["Sex, female / male", f"{(df.SEX=='F').sum()} / {(df.SEX=='M').sum()}"],
        ["Histology", ""],
    ]
    for h,cnt in df.histology.value_counts().items():
        rows.append([f"  {h}", int(cnt)])
    rows += [
        ["Data source", ""],
        ["  Francisco 2022 (fusion+)", int((df.cohort=='Francisco2022').sum())],
        ["  Hirsch 2020 (BAP1)", int((df.cohort=='Hirsch2020').sum())],
    ]
    write_both("Table1_cohort", header, rows)

if __name__ == "__main__":
    table1(); print("DONE Table 1")
